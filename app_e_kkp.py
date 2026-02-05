import streamlit as st
import google.generativeai as genai
import re
from io import BytesIO

# Library Dokumen (Pastikan sudah pip install python-docx fpdf)
from docx import Document
from docx.shared import Pt, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from fpdf import FPDF

# --- 1. KONFIGURASI HALAMAN (WAJIB DI ATAS SENDIRI) ---
st.set_page_config(page_title="KKP Audit Generator", layout="wide")

# --- 2. LIST MODEL ---
AVAILABLE_MODELS = [
    "gemini-1.5-flash", 
    "gemini-1.5-pro", 
    "gemini-2.0-flash-exp",
    "gemini-2.5-flash",
    "gemini-3-flash-preview",
    "gemini-3-pro-preview"
]

# --- 3. PROMPT SYSTEM (Dioptimalkan untuk Parsing) ---
SYSTEM_INSTRUCTION = """
Anda adalah Auditor Senior. Tugas Anda menyusun Kertas Kerja Pemeriksaan (KKP).
PENTING: Ikuti format di bawah ini dengan ketat. Jangan ubah urutan nomor header.

[HEADER_START]
1. No. KKP: [Isi/Strip]
2. Nama Unit Kerja: [Isi]
3. Periode Pemeriksaan: [Isi]
4. INTERNAL AUDITOR: [Isi Nama Tim]
5. AUDITEE: [Isi Nama Auditee]
6. Materi Pemeriksaan: [Isi Judul]
[HEADER_END]

[CONTENT_START]
**URAIAN PEMERIKSAAN**
[PARAGRAPH]
[Isi uraian singkat di sini...]

**CATATAN PEMERIKSA**
[PARAGRAPH]
[Isi temuan detail di sini. Gunakan poin-poin jika perlu...]

**Atas Catatan Pemeriksa, Bahwa Kondisi Tersebut Belum Sesuai Dengan**
[PARAGRAPH]
[Sebutkan peraturan yang dilanggar...]

**Kondisi Tersebut Dapat Mengakibatkan**
[PARAGRAPH]
[Isi dampak...]

**Kondisi Tersebut Disebabkan Oleh**
[PARAGRAPH]
[Isi penyebab...]

**REKOMENDASI**
[PARAGRAPH]
[Isi rekomendasi...]
[CONTENT_END]
"""

# --- 4. FUNGSI GENERATE AI ---
def get_ai_response(api_key, model_name, input_text):
    try:
        genai.configure(api_key=api_key)
        model = genai.GenerativeModel(model_name)
        prompt = f"{SYSTEM_INSTRUCTION}\n\nDATA MENTAH USER:\n{input_text}"
        with st.spinner("🤖 AI sedang menyusun KKP..."):
            response = model.generate_content(prompt)
            return response.text
    except Exception as e:
        st.error(f"Terjadi Kesalahan AI: {str(e)}")
        return None

# --- 5. FUNGSI MEMBUAT WORD (DOCX) - VERSI RAPI (TABLE) ---
def create_docx(ai_text):
    doc = Document()
    
    # Set Font Default Document
    style = doc.styles['Normal']
    style.font.name = 'Arial'
    style.font.size = Pt(11)
    
    # Judul Dokumen
    p_title = doc.add_paragraph()
    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p_title.add_run("PT AGRINAS PANGAN NUSANTARA (PERSERO)\nINTERNAL AUDIT (IA)")
    run.bold = True
    run.font.size = Pt(12)
    
    doc.add_paragraph() # Spasi kosong

    # Parsing Text
    lines = ai_text.split('\n')
    header_mode = False
    
    # Tabel Header (Agar Titik Dua Lurus)
    table = None
    
    for line in lines:
        clean_line = line.strip()
        
        if not clean_line: continue
        if "[HEADER_START]" in clean_line:
            header_mode = True
            # Buat Tabel 3 Kolom: [Label] [:] [Isi]
            table = doc.add_table(rows=0, cols=3)
            table.autofit = False
            # Atur lebar kolom (Label lebar dikit, Colon sempit, Value sisa)
            table.columns[0].width = Cm(5) 
            table.columns[1].width = Cm(0.5)
            table.columns[2].width = Cm(10)
            continue
            
        if "[HEADER_END]" in clean_line:
            header_mode = False
            doc.add_paragraph() # Spasi setelah header
            continue

        if "[CONTENT_START]" in clean_line or "[CONTENT_END]" in clean_line:
            continue

        # LOGIKA HEADER (Masuk ke Tabel)
        if header_mode and ":" in clean_line:
            parts = clean_line.split(":", 1)
            label = parts[0].strip()
            value = parts[1].strip() if len(parts) > 1 else ""
            
            row = table.add_row()
            row.cells[0].text = label
            row.cells[1].text = ":"
            row.cells[2].text = value
            
            # Bold label
            row.cells[0].paragraphs[0].runs[0].bold = True
            
        # LOGIKA KONTEN (Paragraf Biasa)
        elif not header_mode:
            if "[PARAGRAPH]" in clean_line:
                continue # Skip marker
            
            # Deteksi Judul Sub-Bab (Bold)
            if clean_line.startswith("**") or clean_line.isupper():
                p = doc.add_paragraph()
                clean_text = clean_line.replace("**", "")
                run = p.add_run(clean_text)
                run.bold = True
                p.paragraph_format.space_before = Pt(12)
                p.paragraph_format.space_after = Pt(2)
            else:
                # Isi Narasi (Justify)
                p = doc.add_paragraph(clean_line)
                p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    # Simpan ke Buffer
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer

# --- 6. FUNGSI MEMBUAT PDF (FPDF) ---
class PDF(FPDF):
    def header(self):
        self.set_font('Arial', 'B', 12)
        self.cell(0, 5, 'PT AGRINAS PANGAN NUSANTARA (PERSERO)', 0, 1, 'C')
        self.cell(0, 5, 'INTERNAL AUDIT (IA)', 0, 1, 'C')
        self.ln(10)

def create_pdf(ai_text):
    pdf = PDF()
    pdf.add_page()
    pdf.set_auto_page_break(auto=True, margin=15)
    pdf.set_font("Arial", size=11)
    
    lines = ai_text.split('\n')
    header_mode = False
    
    for line in lines:
        clean_line = line.strip()
        if not clean_line: continue
        
        # Filter Marker
        if "[HEADER_START]" in clean_line: header_mode = True; continue
        if "[HEADER_END]" in clean_line: header_mode = False; pdf.ln(5); continue
        if "[CONTENT_START]" in clean_line or "[CONTENT_END]" in clean_line or "[PARAGRAPH]" in clean_line: continue
        
        # Mode Header (Rata Kiri dengan Cell tetap)
        if header_mode and ":" in clean_line:
            parts = clean_line.split(":", 1)
            label = parts[0].strip()
            val = parts[1].strip() if len(parts) > 1 else ""
            
            pdf.set_font("Arial", 'B', 11)
            pdf.cell(50, 6, label, 0, 0, 'L') # Lebar label 50mm
            pdf.set_font("Arial", '', 11)
            pdf.cell(5, 6, ":", 0, 0, 'L')    # Lebar titik dua 5mm
            pdf.multi_cell(0, 6, val, 0, 'L') # Sisanya untuk value
            
        # Mode Konten
        elif not header_mode:
            # Judul Sub (Bold)
            if clean_line.startswith("**") or (len(clean_line) < 50 and clean_line.isupper()):
                pdf.ln(3)
                pdf.set_font("Arial", 'B', 11)
                text = clean_line.replace("**", "")
                pdf.multi_cell(0, 6, text, 0, 'L')
                pdf.set_font("Arial", '', 11)
            else:
                # Narasi (Justify = 'J' terkadang bug di FPDF lama, gunakan 'L' jika error, tapi kita coba 'J')
                # Kita gunakan encode latin-1 agar karakter aneh tidak bikin crash
                try:
                    safe_text = clean_line.encode('latin-1', 'replace').decode('latin-1')
                    pdf.multi_cell(0, 6, safe_text, 0, 'J')
                except:
                    pdf.multi_cell(0, 6, clean_line, 0, 'L') # Fallback ke Left jika error
                
    
    buffer = BytesIO()
    # Output sebagai string byte
    pdf_bytes = pdf.output(dest='S').encode('latin-1')
    buffer.write(pdf_bytes)
    buffer.seek(0)
    return buffer

# --- 7. TAMPILAN UTAMA (UI & EDITOR) ---
st.title("📑 KKP Generator Pro (Editor Mode)")
st.markdown("---")

# Inisialisasi Session State untuk menyimpan draft teks
if 'kkp_draft' not in st.session_state:
    st.session_state.kkp_draft = ""

# Sidebar
with st.sidebar:
    st.header("🔧 Pengaturan")
    api_key = st.text_input("Gemini API Key", type="password")
    
    model_choice = st.selectbox("Pilih Model AI", AVAILABLE_MODELS)
    if model_choice == "Input Manual...":
        model_final = st.text_input("Nama Model Custom", "gemini-1.5-pro")
    else:
        model_final = model_choice
        
    st.info("Status: Siap Digunakan")
    
    st.markdown("---")
    st.header("🖼️ & 📊 Insert Tools")
    
    # Tool Insert Gambar (Simulasi)
    uploaded_file = st.file_uploader("Upload Foto Bukti", type=['png', 'jpg', 'jpeg'])
    if uploaded_file is not None:
        st.caption(f"File: {uploaded_file.name}")
        if st.button("Sisipkan Kode Foto"):
            # Menambahkan tag gambar ke editor
            st.session_state.kkp_draft += f"\n\n[IMAGE: {uploaded_file.name}]\n(Catatan: Pastikan file foto ada di folder yang sama saat convert final)"

# Layout Utama
col_input, col_editor = st.columns([1, 1.2])

with col_input:
    st.subheader("1. Data Temuan Audit")
    raw_data = st.text_area(
        "Paste catatan lapangan di sini:", 
        height=400,
        placeholder="Contoh:\nUnit: Divisi Umum\nTanggal: 5 Feb 2026\nTemuan: Ada selisih kas Rp 500rb..."
    )
    
    generate_btn = st.button("🚀 Buat Draft KKP", type="primary", use_container_width=True)

# Logic Generate AI
if generate_btn:
    if not api_key:
        st.warning("⚠️ Masukkan API Key di sidebar dulu!")
    elif not raw_data:
        st.warning("⚠️ Data temuan masih kosong!")
    else:
        ai_result = get_ai_response(api_key, model_final, raw_data)
        if ai_result:
            st.session_state.kkp_draft = ai_result
            st.rerun() # Refresh halaman untuk menampilkan hasil di editor

# --- BAGIAN EDITOR DENGAN TOMBOL EDIT ---
with col_editor:
    st.subheader("2. Editor & Preview")
    
    # --- TOOLBAR EDITOR ---
    # Baris 1: Alignment
    c1, c2, c3, c4 = st.columns(4)
    if c1.button("⬅️ Left"):
        st.session_state.kkp_draft += "\n[align:left] Teks Rata Kiri [/align]"
    if c2.button("↔️ Center"):
        st.session_state.kkp_draft += "\n[align:center] Teks Tengah [/align]"
    if c3.button("➡️ Right"):
        st.session_state.kkp_draft += "\n[align:right] Teks Rata Kanan [/align]"
    if c4.button("🟰 Justify"):
        st.session_state.kkp_draft += "\n[align:justify] Teks Rata Kanan-Kiri [/align]"

    # Baris 2: Font & Size & Table
    c5, c6, c7, c8 = st.columns([1.5, 1, 1, 1.5])
    
    # Pilihan Font (Hanya menyisipkan tag teks)
    font_name = c5.selectbox("Jenis Font", ["Arial", "Times New Roman", "Calibri"], label_visibility="collapsed")
    font_size = c6.number_input("Size", min_value=8, max_value=24, value=11, label_visibility="collapsed")
    
    if c7.button("🔤 Apply"):
        st.session_state.kkp_draft += f"\n[font:{font_name}][size:{font_size}] Teks Custom [/size][/font]"
        
    if c8.button("📅 Insert Table"):
        table_template = """
\n[TABLE_START]
| Header 1 | Header 2 | Header 3 |
| :--- | :---: | ---: |
| Data 1 | Data 2 | Data 3 |
[TABLE_END]
"""
        st.session_state.kkp_draft += table_template

    # --- TEXT AREA EDITOR (Terikat dengan session_state) ---
    edited_text = st.text_area(
        "Edit Hasil AI di sini:", 
        value=st.session_state.kkp_draft,
        height=500,
        key="editor_area"
    )
    
    # Update session state jika user mengetik manual
    st.session_state.kkp_draft = edited_text
    
    st.info("💡 Tombol di atas akan menambahkan kode format ke bagian paling bawah teks. Silakan Cut/Paste ke posisi yang diinginkan.")

    # Tombol Download
    b1, b2 = st.columns(2)
    
    docx_file = create_docx(st.session_state.kkp_draft)
    with b1:
        st.download_button(
            label="📥 Download Word (.docx)",
            data=docx_file,
            file_name="KKP_Final_Edit.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )
    
    pdf_file = create_pdf(st.session_state.kkp_draft)
    with b2:
        st.download_button(
            label="📥 Download PDF (.pdf)",
            data=pdf_file,
            file_name="KKP_Final_Edit.pdf",
            mime="application/pdf"
        )
